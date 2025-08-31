# ===================== Bar Exam Indexer (Ontario) =====================
# Streamlit app: upload PDF -> exam-style issues w/ pages -> edit -> cheat sheets
# =====================================================================

import os, io, json, re
from typing import List, Dict, Tuple

import streamlit as st
import pandas as pd
import fitz  # PyMuPDF

try:
    import openai
except ImportError:
    openai = None

# ===================== Config =====================
st.set_page_config(page_title="Bar Exam Indexer — Ontario", layout="wide")
st.title("Bar Exam Indexer • Ontario")

# ===================== OpenAI client =====================
api_key = st.secrets.get("OPENAI_API_KEY") or os.getenv("OPENAI_API_KEY")
if not api_key or openai is None:
    client = None
    st.error("No OpenAI API key configured or openai module missing. Add OPENAI_API_KEY in Streamlit Secrets and ensure openai is installed.")
else:
    try:
        openai.api_key = api_key  # For openai-python 1.x and 0.x compatibility
        client = openai
    except Exception as e:
        client = None
        st.error(f"OpenAI client error: {e}")

# ===================== Subject presets & helpers =====================
SUBJECT_PRESETS = {
    "Professional Responsibility": {
        "examples": """
- Conflict of interest — bright-line rule (current vs former client)
- Confidentiality — exceptions (imminent risk; court order)
- Competence & scope — accepting/withdrawing retainer
- Duty of candour to the tribunal — correcting false evidence
- Trust accounting — mixed funds; record-keeping
- Advertising & solicitation — permissible claims; referrals
""",
        "authorities_hint": "LSO Rules of Professional Conduct; By-Laws; discipline cases."
    },
    "Civil": {
        "examples": """
- Appeal route from SCJ to Divisional Court — Civil
- Summary judgment — no genuine issue requiring a trial (RCP r. 20.04)
- Motion to strike vs judgment on pleadings — Rule 21 vs Rule 25
- Security for costs — entitlement and factors (R. 56)
- Interlocutory injunction — RJR-MacDonald test
""",
        "authorities_hint": "Rules of Civil Procedure; Courts of Justice Act; Limitations Act, 2002; ON/CA cases."
    },
    "Criminal": {
        "examples": """
- Search incident to arrest — scope limits
- Detention vs arrest — s. 9 & 10(b) cautions
- Exclusion of evidence — Grant framework (s. 24(2))
- Bail — ladder principle; tertiary ground
- Mens rea for murder vs manslaughter — objective vs subjective
""",
        "authorities_hint": "Criminal Code; Charter; SCC/ONCA cases (Grant, St-Onge Lamoureux, etc.)."
    },
    "Family": {
        "examples": """
- Best interests test — parenting time/decision-making (CLRA/Divorce Act)
- Mobility — Gordon test; material change
- Child support — table amounts vs undue hardship
- Spousal support — entitlement; SSAG ranges
- Division of property — equalization; excluded assets
""",
        "authorities_hint": "Divorce Act; CLRA; FLA; SSAG; ONCA/SCC cases."
    },
    "Public": {
        "examples": """
- Judicial review — prematurity; adequate alternative remedy
- Standard of review — reasonableness vs correctness (Vavilov)
- Procedural fairness — duty trigger & content (Baker factors)
- Delegated legislation — vires; ultra vires analysis
""",
        "authorities_hint": "Vavilov; Baker; SPPA (Ontario); JRPA; Charter."
    },
}

ISSUE_STYLE_EXAMPLES = """
Generic acceptable 'issue' phrasing:
- Appeal route from SCJ to Divisional Court — Civil
- Limitation period — discoverability under s. 5, Limitations Act, 2002
- Rule 21 motion — striking claims with no reasonable cause of action
- Anton Piller order — elements and safeguards
"""

ISSUE_SCHEMA = """Return ONLY valid JSON with this shape:
{
  "issues": [
    {
      "issue": "short exam-style issue statement",
      "triggers": ["2-5 spotting cues"],
      "rule_or_test": "concise elements or test (Ontario)",
      "authorities": ["Rules/statutes/cases"],
      "pages": [int, ...]
    }
  ]
}"""

def subject_preset_block(subject: str) -> str:
    p = SUBJECT_PRESETS.get(subject, {})
    out = ""
    if p.get("examples"):
        out += f"\nSUBJECT-SPECIFIC EXAMPLES ({subject}):\n{p['examples']}\n"
    if p.get("authorities_hint"):
        out += f"\nWhen citing authorities, prefer: {p['authorities_hint']}\n"
    return out

def _granularity_text(level: str) -> str:
    if level.startswith("Fine"):
        return "Write concrete exam issues (route/test/standard/duty/remedy), not headings."
    if level.startswith("Medium"):
        return "Write practical sub-issues useful for exam answers."
    return "High level issues; fewer items."

def build_chunk_prompt(joined_pages: str, granularity: str, subject_hint: str, require_auth: bool, subject_preset_name: str) -> str:
    need_auth = ("Include specific Ontario/Canadian authorities." if require_auth else "Do not include citations unless essential.")
    extraction_policy = (
        "You are an assistant helping a bar exam student. Extract ALL POSSIBLE LEGAL ISSUES, no matter how small. "
        "If the text is mainly instructive or TOC, infer procedural/context issues instead of saying 'none'."
    )
    preset = subject_preset_block(subject_preset_name)
    return f"""You are building an exam index for the Ontario bar ({subject_hint}).

{_granularity_text(granularity)}
Avoid TOC-style headings. Each 'issue' must read like an exam spotter, e.g., 'Appeal route from SCJ to Divisional Court — Civil'.

EXAMPLES (generic):
{ISSUE_STYLE_EXAMPLES}
{preset}

{extraction_policy}

From the following pages, extract 12–25 issues using the JSON schema. {need_auth}

{ISSUE_SCHEMA}

Pages:
{joined_pages}
"""

# -------- Cheat sheet prompt builder (subject-aware) --------
CHEAT_GENERIC_STYLE = """
Write compact exam cheat sheets. For each issue:
- Definition (1–2 lines)
- Elements / test (numbered, concise)
- Pitfalls / exceptions (bullets)
- Tiny checklist (3–5 checks)
Keep each under ~120 words. Use Ontario framing.
"""

CHEAT_SUBJECT_TIPS = {
    "Professional Responsibility": """
- Emphasize LSO Rules of Professional Conduct; by-laws; commentary.
- Include classic pitfalls: conflicts, confidentiality exceptions, candour, trust accounting, advertising, civility, supervision.
- Use verbs like "must", "may", "avoid", "disclose", "withdraw".
""",
    "Civil": """
- Prefer Rules of Civil Procedure cites (rule numbers) and CJA/Limitations Act hooks.
- Emphasize tests (RJR-MacDonald, Hryniak), timelines, burdens, standards of review, motion types.
""",
    "Criminal": """
- Anchor to Criminal Code/Charter. Stress Grant (24(2)), detention/arrest, right to counsel, voluntariness, warrants/exceptions, bail ladder.
""",
    "Family": """
- Prefer Divorce Act/CLRA/FLA + SSAG. Emphasize best interests, mobility (Gordon), support (entitlement/quantum), equalization nuances, domestic contract validity.
""",
    "Public": """
- Emphasize Vavilov/Baker frameworks: standard of review, duty of fairness (trigger & content), judicial review remedies, vires/ultra vires.
""",
}

def build_cheat_prompt(
    selected_rows: List[Dict],
    subject_name: str,
    subject_hint: str,
    include_cases: bool,
) -> str:
    issues_block = "\n".join(
        "- " + (r.get("issue","").strip())
        + (f"\n  Rule/Test: {r['rule_or_test']}" if r.get("rule_or_test") else "")
        + (f"\n  Authorities: {r['authorities']}" if r.get("authorities") else "")
        + (f"\n  Pages: {r['pages']}" if r.get("pages") else "")
        for r in selected_rows if r.get("issue")
    )
    cases_hint = (
        "When helpful, include 1–2 key Ontario/Canadian cases or rule/statute cites (super short parentheticals)."
        if include_cases else
        "Do not include case citations; focus on elements/tests."
    )
    subject_tips = CHEAT_SUBJECT_TIPS.get(subject_name, "")
    return f"""Create concise Ontario bar exam cheat sheets for the following **{subject_hint}** issues.

{CHEAT_GENERIC_STYLE}
Subject tips:
{subject_tips}
{cases_hint}

Issues:
{issues_block}

Return GitHub-flavored Markdown with ### headings for each issue name.
Keep language exam-ready and precise. Do not exceed ~120 words per issue.
"""

# ===================== Sidebar =====================
with st.sidebar:
    st.header("Settings")
    # Only show actual available models
    model = st.selectbox("Model", ["gpt-4o", "gpt-4-turbo", "gpt-3.5-turbo"], index=0)
    temperature = st.slider("Temperature", 0.0, 2.0, 0.2, 0.05)
    chunk_size = st.number_input("Pages per chunk", min_value=2, max_value=25, value=8, step=1)
    max_issues = st.number_input("Max issues to keep", min_value=10, max_value=200, value=80, step=5)

    st.header("Extraction style")
    subject_preset = st.selectbox("Subject preset", list(SUBJECT_PRESETS.keys()), index=0)
    subject_hint = st.text_input("Subject focus (optional override)", value=subject_preset)
    granularity = st.selectbox("Granularity", ["Fine (exam spotting)", "Medium", "Coarse"], index=0)
    require_authorities = st.checkbox("Include statutes/rules/cases", value=True)

    st.header("Debug")
    debug_show_raw = st.checkbox("Show raw model output per chunk", value=False)

# ===================== Upload PDF (with fallback + preview) =====================
if "pages_text" not in st.session_state:
    st.session_state["pages_text"] = []

uploaded = st.file_uploader("Upload PDF (<= 200 MB)", type=["pdf"])

if uploaded:
    # Prompt for original page numbers right after upload
    start_page = st.number_input(
        "Enter the starting page number in the original materials:",
        min_value=1, value=1, step=1
    )
    end_page = st.number_input(
        "Enter the ending page number in the original materials:",
        min_value=start_page, value=start_page, step=1
    )
    # Now you can use start_page and end_page in your processing
    # ... rest of your PDF extraction/preview logic ...

def extract_with_pymupdf(pdf_bytes: bytes) -> List[Tuple[int, str]]:
    out: List[Tuple[int, str]] = []
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    for i, page in enumerate(doc):
        try:
            out.append((i + 1, page.get_text("text") or ""))
        except Exception:
            out.append((i + 1, ""))
    doc.close()
    return out

def extract_with_pdfminer_if_available(pdf_bytes: bytes, page_numbers: List[int]) -> Dict[int, str]:
    """Try pdfminer fallback; if unavailable, return {} without raising."""
    try:
        from pdfminer.high_level import extract_text  # requires pdfminer.six
    except Exception:
        return {}
    tmp = io.BytesIO(pdf_bytes)
    out: Dict[int, str] = {}
    for p in page_numbers:
        try:
            tmp.seek(0)
            out[p] = extract_text(tmp, page_numbers=[p - 1]) or ""
        except Exception as e:
            st.warning(f"pdfminer error on page {p}: {e}")
            out[p] = ""
    return out

if uploaded is not None:
    try:
        pdf_bytes = uploaded.getvalue()
        raw = extract_with_pymupdf(pdf_bytes)
        empty_pages = [p for (p, t) in raw if not (t and t.strip())]
        if empty_pages:
            pm_texts = extract_with_pdfminer_if_available(pdf_bytes, empty_pages)
            if pm_texts:
                fixed: List[Tuple[int, str]] = []
                for (p, t) in raw:
                    if (not t or not t.strip()) and pm_texts.get(p):
                        t = pm_texts[p]
                    fixed.append((p, t))
                raw = fixed
            else:
                st.info("pdfminer.six not installed or fallback unavailable. If many pages are blank, run OCR and re-upload.")

        # Assume user_start_page is the value the user entered in your input box
st.session_state["pages_text"] = [f"[Page {user_start_page + i}]\n{t}" for i, (p, t) in enumerate(raw)]
        total_chars = sum(len(t) for (_, t) in raw)
        st.success(f"Loaded {len(st.session_state['pages_text'])} pages from {uploaded.name} • {total_chars} characters extracted")
        preview_n = min(2, len(st.session_state["pages_text"]))
        st.caption("Preview (to confirm we actually have text):")
        st.text("\n\n".join(st.session_state["pages_text"][:preview_n])[:3000] or "(no extractable text)")
        empty_cnt = sum(1 for (_, t) in raw if not (t and t.strip()))
        if empty_cnt >= max(1, int(len(raw) * 0.7)):
            st.warning("Most pages had no extractable text. Your PDF may be scanned images. Run OCR (e.g., Acrobat 'Recognize Text') and re-upload.")
    except Exception as e:
        st.error(f"Error reading PDF: {e}")

# ===================== Step 1 • Generate exam-style issues (with page refs) =====================
st.subheader("Step 1 • Generate exam-style issues (with page refs)")

def parse_rows_from_json(txt: str) -> List[Dict]:
    try:
        # Remove ```json or ``` if present
        txt = re.sub(r"^```json|^```|```$", "", txt.strip(), flags=re.MULTILINE)
        # Extract JSON block if mixed with prose
        match = re.search(r"(\{[\s\S]+\})", txt)
        if match:
            txt = match.group(1)
        data = json.loads(txt)
        rows = data.get("issues", [])
        if isinstance(rows, list):
            normed = []
            for it in rows:
                normed.append({
                    "issue": (it.get("issue") or "").strip(),
                    "triggers": "; ".join(it.get("triggers") or []),
                    "rule_or_test": (it.get("rule_or_test") or "").strip(),
                    "authorities": "; ".join(it.get("authorities") or []),
                    "pages": ", ".join(str(p) for p in (it.get("pages") or [])),
                })
            return [r for r in normed if r["issue"]]
    except Exception:
        pass
    return []

if st.button("Generate issues"):
    pages_text = st.session_state.get("pages_text", [])
    if not pages_text:
        st.warning("Please upload a PDF first (and confirm the preview shows text).")
    elif not client:
        st.warning("OpenAI client not configured.")
    else:
        chunks = [pages_text[i:i + int(chunk_size)] for i in range(0, len(pages_text), int(chunk_size))]
        all_rows: List[Dict] = []

        for idx, ch in enumerate(chunks, start=1):
            joined = "\n\n".join(ch)
            prompt = build_chunk_prompt(joined, granularity, subject_hint, require_authorities, subject_preset)

            def show_raw(label, content):
                if debug_show_raw:
                    with st.expander(f"Chunk {idx} • {label} (raw)"):
                        st.text(content[:8000])

            try:
                resp = client.chat.completions.create(
                    model=model,
                    temperature=temperature,
                    messages=[
                        {"role": "system", "content": "You extract exam issues for the Ontario bar. Be precise and practical. Return JSON."},
                        {"role": "user", "content": prompt},
                    ],
                    max_tokens=1500,
                )
                out = resp.choices[0].message.content.strip()
                show_raw("pass 1", out)
                rows = parse_rows_from_json(out)

                # Pass 2: strict JSON retry
                if not rows:
                    resp2 = client.chat.completions.create(
                        model=model,
                        temperature=temperature,
                        messages=[
                            {"role": "system", "content": "Return ONLY the JSON object that matches the schema. No prose."},
                            {"role": "user", "content": prompt + "\n\nReturn ONLY the JSON object, nothing else."},
                        ],
                        max_tokens=1500,
                    )
                    out2 = resp2.choices[0].message.content.strip()
                    show_raw("pass 2 (strict JSON)", out2)
                    rows = parse_rows_from_json(out2)

                # Pass 3: never-empty fallback
                if not rows:
                    fallback_prompt = prompt + """
IMPORTANT FALLBACK:
If the content is a cover, preface, overview, or TOC, infer procedural/contextual issues a candidate needs from that section
(e.g., motion types & tests (R. 20, R. 21), appeal routes/standards, limitation/discoverability, filing/service timelines, evidentiary burdens, remedies).
Return AT LEAST 10 items in the same JSON schema (never an empty list).
"""
                    resp3 = client.chat.completions.create(
                        model=model,
                        temperature=temperature,
                        messages=[
                            {"role": "system", "content": "Return ONLY the JSON object that matches the schema. No prose."},
                            {"role": "user", "content": fallback_prompt + "\n\nReturn ONLY the JSON object, nothing else."},
                        ],
                        max_tokens=1500,
                    )
                    out3 = resp3.choices[0].message.content.strip()
                    show_raw("pass 3 (forced fallback)", out3)
                    rows = parse_rows_from_json(out3)

                if not rows:
                    st.warning(f"Chunk {idx}: no issues parsed after 3 passes.")
                else:
                    all_rows.extend(rows)

            except Exception as e:
                st.error(f"OpenAI error on chunk {idx}: {e}")
                break

        st.session_state["issues_rows"] = all_rows[: int(max_issues)]
        if st.session_state["issues_rows"]:
            st.success(f"Extracted {len(st.session_state['issues_rows'])} issues.")
        else:
            st.error("No issues were extracted. Check the preview above — if it’s mostly blank, run OCR on your PDF and try again.")

# ===================== Step 2 • Review and edit the index =====================
st.subheader("Step 2 • Review and edit the index")

if st.session_state.get("issues_rows"):
    df = pd.DataFrame(st.session_state["issues_rows"])
    edited = st.data_editor(
        df,
        num_rows="dynamic",
        use_container_width=True,
        column_config={
            "issue": st.column_config.TextColumn("Issue (exam-style)", width="medium"),
            "triggers": st.column_config.TextColumn("Spotting triggers", width="medium"),
            "rule_or_test": st.column_config.TextColumn("Rule / test (concise)", width="large"),
            "authorities": st.column_config.TextColumn("Authorities (rules/statutes/cases)", width="large"),
            "pages": st.column_config.TextColumn("Pages", width="small"),
        },
        key="data_editor_issues",
    )
    st.session_state["issues_rows"] = edited.to_dict("records")
    c1, c2 = st.columns(2)
    with c1:
        st.download_button(
            "Download index as CSV",
            pd.DataFrame(st.session_state["issues_rows"]).to_csv(index=False),
            "bar_index.csv",
            mime="text/csv",
        )
    with c2:
        st.download_button(
            "Download index as Markdown",
            pd.DataFrame(st.session_state["issues_rows"]).to_markdown(index=False),
            "bar_index.md",
            mime="text/markdown",
        )

# ===================== Step 3 • Generate cheat sheets (subject-aware) =====================
st.subheader("Step 3 • Generate cheat sheets for issues")

issues = st.session_state.get("issues_rows", [])
n_issues = len(issues)

if n_issues == 0:
    st.info("Create or import an index first (Steps 1–2).")
else:
    max_allowed = min(50, n_issues)
    default_val = min(10, max_allowed)

    c1, c2, c3 = st.columns([1, 1, 2])
    with c1:
        top_k = st.number_input(
            "Number of issues to include",
            min_value=1, max_value=max_allowed, value=default_val, step=1,
        )
    with c2:
        include_cases = st.checkbox("Include cases / rule cites", value=True)
    with c3:
        st.caption(f"Cheat sheets are tailored for **{subject_preset}** (set in the sidebar).")

    if st.button("Generate cheat sheets"):
        if client is None:
            st.warning("OpenAI client not configured.")
        else:
            selected = issues[: int(top_k)]
            cheat_prompt = build_cheat_prompt(
                selected_rows=selected,
                subject_name=subject_preset,
                subject_hint=subject_hint,
                include_cases=include_cases,
            )
            try:
                resp = client.chat.completions.create(
                    model=model,
                    temperature=temperature,
                    messages=[
                        {"role": "system", "content": "You produce compact, accurate bar exam cheat sheets."},
                        {"role": "user", "content": cheat_prompt},
                    ],
                    max_tokens=1800,
                )
                md = resp.choices[0].message.content.strip()
                st.session_state["cheat_sheets_md"] = md
                st.success("Cheat sheets created")
            except Exception as e:
                st.error(f"OpenAI error while creating cheat sheets: {e}")

if st.session_state.get("cheat_sheets_md"):
    st.markdown("### Cheat sheets (editable)")
    cheat_md = st.text_area("Markdown", st.session_state["cheat_sheets_md"], height=500, key="cheat_md_area")
    st.session_state["cheat_sheets_md"] = cheat_md
    st.download_button(
        "Download cheat sheets as Markdown",
        st.session_state["cheat_sheets_md"],
        "cheat_sheets.md",
        mime="text/markdown",
    )
    # --- Enhanced Chart Mode ---
    st.markdown("---")
    st.markdown("#### Enhanced Chart Mode (Optional)")
    enhanced_chart_mode = st.checkbox("Enable Enhanced Chart Mode", value=False)
    if enhanced_chart_mode:
        if st.button("Create Chart"):
            if client is None:
                st.warning("OpenAI client not configured.")
            else:
                chart_prompt = (
                    "Take this cheat sheet and create an enhanced Cheat Sheet. "
                    "1) include key legal principles, cases, statutes, and concise notes "
                    "2) organize clearly for quick study and review "
                    "3) keep bullet-point style, readable and compact. "
                    f"Cheat Sheet Input:\n{st.session_state['cheat_sheets_md']}"
                )
                try:
                    with st.spinner("Creating enhanced chart..."):
                        resp = client.chat.completions.create(
                            model=model,
                            temperature=temperature,
                            messages=[
                                {"role": "system", "content": "You are a legal study assistant producing highly organized, chart-style cheat sheets."},
                                {"role": "user", "content": chart_prompt},
                            ],
                            max_tokens=1800,
                        )
                        enhanced_chart = resp.choices[0].message.content.strip()
                        st.session_state["enhanced_chart"] = enhanced_chart
                        st.success("Enhanced chart created!")
                except Exception as e:
                    st.error(f"OpenAI error while creating enhanced chart: {e}")

        # Display the enhanced chart
        if st.session_state.get("enhanced_chart"):
            st.markdown("#### Enhanced Chart Output")
            st.markdown(st.session_state["enhanced_chart"])
            st.download_button(
                "Download Enhanced Chart as Markdown",
                st.session_state["enhanced_chart"],
                "enhanced_chart.md",
                mime="text/markdown",
            )
import io
from docx import Document

def markdown_to_docx(md_text):
    """Simple function to convert markdown to docx. Handles headings, bold, italics, lists, and plain text."""
    doc = Document()
    lines = md_text.splitlines()
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph("")  # Blank line
        elif line.startswith("### "):  # H3
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):  # H2
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):  # H1
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- "):  # Bullet list
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith("1. "):  # Numbered list
            doc.add_paragraph(line[3:], style='List Number')
        else:
            doc.add_paragraph(line)
    return doc

# Inside your Enhanced Chart Mode display section:
if st.session_state.get("enhanced_chart"):
    st.markdown("#### Enhanced Chart Output")
    st.markdown(st.session_state["enhanced_chart"])

    # DOCX download
    doc = markdown_to_docx(st.session_state["enhanced_chart"])
    docx_io = io.BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)
    st.download_button(
        "Download Enhanced Chart as Word (.docx)",
        data=docx_io,
        file_name="enhanced_chart.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="enhanced_chart_docx_btn"
    )

    # Markdown download (optional: keep for reference)
    st.download_button(
        "Download Enhanced Chart as Markdown",
        st.session_state["enhanced_chart"],
        "enhanced_chart.md",
        mime="text/markdown",
        key="enhanced_chart_md_btn"
    )
import streamlit as st
import io
from docx import Document

def generate_study_notes_prompt(issue_name, rule, authorities):
    prompt = (
        f"Write concise, exam-focused study notes for the following bar exam topic.\n"
        f"Issue: {issue_name}\n"
        f"Rule/Test: {rule}\n"
        f"Authorities: {authorities}\n"
        "Include key principles, rationale, important cases/statutes, and tips for exam success. Use clear bullet points."
    )
    return prompt

def notes_to_docx(notes_dict):
    doc = Document()
    doc.add_heading("Bar Exam Study Notes", level=0)
    for issue, notes in notes_dict.items():
        doc.add_heading(issue, level=1)
        doc.add_paragraph(notes)
    return doc

# Main Study Notes Section
if st.session_state.get("issues_rows"):
    st.markdown("## Study Notes")
    notes_dict = st.session_state.get("study_notes_dict", {})
    for idx, issue_row in enumerate(st.session_state["issues_rows"]):
        issue = issue_row.get("issue", f"Issue {idx+1}")
        rule = issue_row.get("rule_or_test", "")
        authorities = issue_row.get("authorities", "")

        st.markdown(f"**{issue}**")
        default_notes = notes_dict.get(issue, "")
        notes = st.text_area(f"Study Notes for {issue}", value=default_notes, key=f"notes_{idx}")

        # Generate button
        if st.button(f"Generate Study Notes for {issue}", key=f"gen_notes_{idx}"):
            if client is None:
                st.warning("OpenAI client not configured.")
            else:
                prompt = generate_study_notes_prompt(issue, rule, authorities)
                with st.spinner("Generating notes..."):
                    resp = client.chat.completions.create(
                        model=model,
                        temperature=temperature,
                        messages=[
                            {"role": "system", "content": "You are a legal bar exam coach who writes smart, concise, practical study notes."},
                            {"role": "user", "content": prompt},
                        ],
                        max_tokens=500,
                    )
                    notes = resp.choices[0].message.content.strip()
                    st.session_state.setdefault("study_notes_dict", {})[issue] = notes
                    st.success("Study notes generated!")

        # Save edited notes
        st.session_state.setdefault("study_notes_dict", {})[issue] = notes

    # Download all notes as Word
    if st.session_state.get("study_notes_dict"):
        doc = notes_to_docx(st.session_state["study_notes_dict"])
        docx_io = io.BytesIO()
        doc.save(docx_io)
        docx_io.seek(0)
        st.download_button(
            "Download All Study Notes as Word (.docx)",
            data=docx_io,
            file_name="study_notes.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="study_notes_docx_btn"
        )
